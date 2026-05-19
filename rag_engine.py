from langchain_ollama import ChatOllama, OllamaEmbeddings
from langchain_community.vectorstores import FAISS
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.output_parsers import StrOutputParser
from langchain_core.runnables import RunnablePassthrough
import os
from pypdf import PdfReader
from docx import Document as DocxDocument

# 初始化本地大模型
llm = ChatOllama(
    model="qwen2:7b-instruct-q4_0",
    temperature=0.1,
    base_url="http://localhost:11434",
    streaming=True
)

# 初始化本地嵌入模型
embeddings = OllamaEmbeddings(
    model="bge-m3",
    base_url="http://localhost:11434"
)

# 文档解析函数（优化了PDF解析，处理大文件更稳定）
def load_document(file_path):
    """加载并解析不同格式的文档"""
    ext = os.path.splitext(file_path)[1].lower()
    text = ""
    
    if ext == ".pdf":
        reader = PdfReader(file_path)
        total_pages = len(reader.pages)
        print(f"正在解析PDF，共{total_pages}页...")
        for i, page in enumerate(reader.pages):
            if i % 50 == 0:
                print(f"已解析{i}/{total_pages}页")
            text += page.extract_text() + "\n"
        print("PDF解析完成！")
    elif ext == ".docx":
        doc = DocxDocument(file_path)
        for para in doc.paragraphs:
            text += para.text + "\n"
    elif ext in [".txt", ".md"]:
        with open(file_path, "r", encoding="utf-8") as f:
            text = f.read()
    else:
        raise ValueError(f"不支持的文件格式: {ext}")
    
    return [Document(page_content=text, metadata={"source": os.path.basename(file_path)})]

# 构建RAG链
def build_rag_chain(documents):
    """从文档构建完整的RAG问答链"""
    # 1. 文本分块（优化了分块参数，更适合产品手册）
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=800,
        chunk_overlap=150,
        separators=["\n\n", "\n", ".", "。", " ", ""]
    )
    splits = text_splitter.split_documents(documents)
    print(f"文档分块完成，共{len(splits)}个块")
    
    # 2. 创建FAISS向量数据库（速度最快，最稳定）
    print("正在构建向量数据库...")
    vectorstore = FAISS.from_documents(
        documents=splits,
        embedding=embeddings
    )
    print("向量数据库构建完成！")
    
    # 3. 创建检索器
    retriever = vectorstore.as_retriever(
        search_type="mmr",
        search_kwargs={"k": 4}
    )
    
    # 4. 定义Prompt模板（优化了产品手册问答）
    prompt = ChatPromptTemplate.from_template("""
你是一个专业的产品文档助手。请严格基于以下提供的产品手册内容回答用户问题。

回答规则：
1. 只使用上下文中明确提到的信息，绝对不要编造内容
2. 如果上下文中没有相关信息，直接回答"抱歉，我在产品手册中没有找到相关信息"
3. 回答要简洁、准确、有条理，分点说明复杂问题
4. 引用手册中的原文时用引号标注
5. 不要提及"上下文"、"手册"等提示词，直接给出答案

上下文：
{context}

用户问题：{question}

专业回答：
""")
    
    # 5. 构建LCEL链
    def format_docs(docs):
        return "\n\n".join(doc.page_content for doc in docs)
    
    rag_chain = (
        {"context": retriever | format_docs, "question": RunnablePassthrough()}
        | prompt
        | llm
        | StrOutputParser()
    )
    
    return rag_chain, vectorstore